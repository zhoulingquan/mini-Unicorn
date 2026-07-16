#!/usr/bin/env python3
"""用 python-docx 创建 Word 文档。

用法:
    # 基本文档
    python create_docx.py output.docx --title "报告标题" --paragraphs "第一段" "第二段"

    # 多级标题 + 段落
    python create_docx.py output.docx --title "报告" --heading2 "第一章" --paragraphs "内容1" --heading2 "第二章" --paragraphs "内容2"

    # 列表
    python create_docx.py output.docx --bullets "项目A" "项目B" "项目C"
    python create_docx.py output.docx --numbers "步骤1" "步骤2"

    # 表格
    python create_docx.py output.docx --table '[["姓名","年龄"],["张三",25],["李四",30]]'

    # 分页
    python create_docx.py output.docx --title "第一页" --page-break --paragraphs "第二页内容"

    # 页眉页脚
    python create_docx.py output.docx --header "公司名称" --footer "第 {page} 页"

    # 组合用法
    python create_docx.py output.docx \\
        --title "季度报告" \\
        --heading2 "Q1 业绩" \\
        --paragraphs "本季度营收同比增长 15%。" \\
        --table '[["指标","数值"],["营收","1.2亿"],["增长率","15%"]]' \\
        --heading2 "Q2 展望" \\
        --bullets "拓展新市场" "优化产品线"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def create_docx(
    output_path: Path,
    *,
    title: str | None = None,
    paragraphs: list[str] | None = None,
    heading2: list[str] | None = None,
    bullets: list[str] | None = None,
    numbers: list[str] | None = None,
    table: str | None = None,
    page_breaks: bool = False,
    header: str | None = None,
    footer: str | None = None,
) -> None:
    """创建 .docx 文档。"""
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Pt
    except ImportError:
        print("错误:需要 python-docx 库。安装:pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # 设置默认字体(中文友好)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # 页眉
    if header:
        for section in doc.sections:
            section.header.paragraphs[0].text = header

    # 页脚(支持 {page} 占位符)
    if footer:
        for section in doc.sections:
            p = section.footer.paragraphs[0]
            text = footer.replace("{page}", "")
            p.text = text
            # python-docx 不直接支持页码字段,这里加文字提示
            if "{page}" in footer:
                run = p.add_run(" (页码字段需手动添加)")
                run.italic = True

    # 标题
    if title:
        doc.add_heading(title, level=0)

    # heading2 + paragraphs 按顺序处理
    # 支持交替使用:--heading2 A --paragraphs p1 --heading2 B --paragraphs p2
    # 这里简化为:所有 heading2 先加,所有 paragraphs 后加
    # 复杂场景请用 unpack/pack 流程
    if heading2:
        for h in heading2:
            doc.add_heading(h, level=1)

    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)

    # 项目符号列表
    if bullets:
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    # 编号列表
    if numbers:
        for item in numbers:
            doc.add_paragraph(item, style="List Number")

    # 表格
    if table:
        try:
            rows_data = json.loads(table)
        except json.JSONDecodeError as e:
            print(f"错误:--table JSON 解析失败: {e}", file=sys.stderr)
            sys.exit(1)
        if not rows_data or not isinstance(rows_data, list):
            print("错误:--table 必须是二维数组", file=sys.stderr)
            sys.exit(1)
        rows = len(rows_data)
        cols = max(len(r) for r in rows_data)
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = "Table Grid"
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    tbl.cell(i, j).text = str(cell_text)

    # 分页
    if page_breaks:
        doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已创建: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="创建 Word 文档")
    parser.add_argument("output", help="输出 .docx 路径")
    parser.add_argument("--title", help="主标题(H1)")
    parser.add_argument("--paragraphs", nargs="+", help="段落文本(可多个)")
    parser.add_argument("--heading2", nargs="+", help="二级标题(可多个)")
    parser.add_argument("--bullets", nargs="+", help="项目符号列表")
    parser.add_argument("--numbers", nargs="+", help="编号列表")
    parser.add_argument("--table", help='表格 JSON,如 \'[["A","B"],["1","2"]]\'')
    parser.add_argument("--page-break", action="store_true", help="末尾添加分页")
    parser.add_argument("--header", help="页眉文字")
    parser.add_argument("--footer", help="页脚文字({page} 占位符)")

    args = parser.parse_args()
    create_docx(
        Path(args.output).resolve(),
        title=args.title,
        paragraphs=args.paragraphs,
        heading2=args.heading2,
        bullets=args.bullets,
        numbers=args.numbers,
        table=args.table,
        page_breaks=args.page_break,
        header=args.header,
        footer=args.footer,
    )


if __name__ == "__main__":
    main()
